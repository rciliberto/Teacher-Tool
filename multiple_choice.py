"""
Includes multiple choice test and question classes and related final static variables
"""
from docx import Document

class MCAnswer(object):
    """An answer object thet stores an answer string ans whether it is true or not"""

    def __init__(self, answer, is_correct):
        self.answer = answer
        self.is_correct = is_correct

    def __str__(self):
        return self.answer

class MCQuestion(object):
    """A question object that stores a question and its answers"""
    # string whose indexes are accessed for multiple choices
    LETTERS = "abcdefg"

    # create a question string, an answer array, and an index for the
    # correct answer in the answer array
    def __init__(self):
        self.question = ""
        self.answers = []

    # print the question followed by the answers.
    # The letter is determined based on the index of the answer in the answer array.
    def __str__(self):
        output = "  " + self.question + "\n"
        for i in range(len(self.answers)):
            output += "     " + MCQuestion.LETTERS[i] + ". " + str(self.answers[i]) + "\n"
        return output

    def set_question(self, question):
        """sets the question"""
        self.question = str(question)

    # add answers to the answer array, and if it's correct,
    # set correct_index to the correct answer's index
    def add_answer(self, answer, is_correct):
        """add an answer to the question"""
        self.answers.append(MCAnswer(answer, is_correct))

class MCTest(object):
    """A test object that stores multiple MCQuestion objects"""

    # set the test name and create a questions array
    def __init__(self):
        self.test_name = ""
        self.questions = []

    # print the title, followed by each question
    def __str__(self):
        output = self.test_name + "\n"
        for question in self.questions:
            output += str(question) + "\n"
        return output

    def add_question(self, question):
        """add a question to the test"""
        self.questions.append(question)

    def set_name(self, name):
        """Set the name of the test"""
        self.test_name = str(name)

    def make_docx(self, doc_name):
        """Make a .docx document representation of the test"""
        # create a doc, and set heading name
        doc = Document()
        doc.add_heading(self.test_name, 0)

        # cycle through each question to add it to the doc
        for i in range(len(self.questions)):
            doc.add_paragraph("\n    " + str(i+1) + ". " + self.questions[i].question)

            # cycle through each answer in the question to add it to the question
            for j in range(len(self.questions[i].answers)):
                doc.add_paragraph("        " + MCQuestion.LETTERS[j] + ". "\
                    "" + str(self.questions[i].answers[j]) + "")

        # save the document
        doc.add_paragraph("")
        doc.save(doc_name + ".docx")

    def make_answer_key(self):
        """Make an answer key"""
        doc = Document()
        doc.add_heading(self.test_name + " Answers", 0)

        for i in range(len(self.questions)):
            question = self.questions[i]

            doc.add_paragraph("\n    " + str(i+1) + ". " + question.question)
            for j in range(len(question.answers)):
                if question.answers[j].is_correct:
                    doc.add_paragraph("        " + MCQuestion.LETTERS[j] + ". "\
                        "" + str(question.answers[j]))

        doc.add_paragraph("")
        doc.save(self.test_name + "_ans.docx")
