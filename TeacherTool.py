"""
Teacher Tool is a suite of useful tools for teachers.
These include a test generator at the moment.
"""
import os
import random
from docx import Document

from multiple_choice import MCQuestion, MCTest


def clear():
    """Clear terminal window depending on OS"""
    # windows
    if os.name == "nt":
        os.system("cls")
    # UNIX/MacOS
    elif os.name == "posix":
        os.system("clear")

def choice_creator():
    """Ask for questions & answers and print the generated test"""
    # creates a test document
    doc = Document()
    test = MCTest()

    # ask for a test name and add it to the doc
    test_name = input("What is the name of your test?\n> ")
    test.set_name(test_name)
    doc.add_heading(test.test_name, 0)

    # ask for number of questions to be filled out
    number_questions = int(input("\nHow many multiple choice questions do you want?\n> "))

    # ask for the number of answers per question
    choices_per_question = int(input("\nHow many choices per question? (Max 7)\n> "))
    #test = ""

    for i in range(number_questions):
        question = MCQuestion()

        quest = input("\nWhat is question number " + str(i+1) + "?\n> ")
        question.set_question(quest)
        doc.add_paragraph(str(i+1) + ". " + quest)

        corr = input("\nWhat is the correct answer?\n> ")
        question.add_answer(corr, True)

        for j in range(choices_per_question-1):
            ans = input("\nType a possible answer.\n> ")
            question.add_answer(ans, False)
            doc.add_paragraph("  " + MCQuestion.LETTERS[j] + ". " + ans)
        doc.add_paragraph("")

    doc.save(test_name + ".docx")
    print("Sucessfully Created test!")
    print(test)

def scrambler(test):
    """scrambles both the questions and the answers for each question when a test is given"""
    newtest = MCTest()
    for i in range(len(test.questions)):
        newtest.questions[i] = MCQuestion()

    for i in range(len(test.questions)):
        index = random.randint(0, len(test.questions) - 1)
        while newtest.questions[index].question != "":
            index = random.randint(0, len(test.questions) - 1)
        newtest.questions[index] = test.questions[i]

        for j in range(len(test.questions[i].answers)):
            ansindex = random.randint(0, len(test.questions[i].answers) - 1)
            while newtest.questions[index].answers[ansindex] != "":
                ansindex = random.randint(0, len(test.questions[i].answers) - 1)
            newtest.questions[i].answers[ansindex] = test.questions[i].answers[j]

    return newtest

"""
MAIN
"""

clear()
# introduce the program
print("Hello! Welcome to TeacherTool. This is your one stop shop for all your teaching needs.\n"\
"type \"help\" for directions and possible commands.")

#Begin inputs
while True:
    USER_CHOICE = str(input("> "))

    if USER_CHOICE.upper() == "EXIT":
        break
    elif USER_CHOICE.upper() == "HELP":
        print("\nList of possible commands: \n"\
        "CLEAR      -   Clear the screen\n"\
        "EXIT       -   Exit TeacherTool\n"\
        "HELP       -   Present this list of commands\n"\
        "\n"\
        "1          -   Make a multiple choice test\"\n")

    elif USER_CHOICE.upper() == "CLEAR":
        clear()
    elif USER_CHOICE.upper() == "1":
        choice_creator()

    else:
        print("Unknown command. type \"help\" for a list of commands")
